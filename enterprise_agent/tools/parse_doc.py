from pathlib import Path
from typing import Iterable


_DOC_TYPES = {
    "policies": "policy",
    "workflows": "workflow",
    "projects": "project",
    "meetings": "meeting",
    "contracts": "contract",
    "reports": "report",
}


def parse_document(path: str) -> dict:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix not in {".md", ".txt", ".pdf", ".docx", ".html", ".htm"}:
        raise ValueError("Only .md, .txt, .pdf, .docx, and .html documents are supported")

    content = _read_content(file_path).strip()
    front_matter, content = _split_front_matter(content)
    metadata = {"path": str(file_path), "raw_format": _raw_format(file_path)}
    metadata.update(front_matter)
    return {
        "source": file_path.name,
        "doc_type": _detect_doc_type(file_path),
        "title": _extract_title(file_path, content),
        "content": content,
        "metadata": metadata,
    }


def _detect_doc_type(path: Path) -> str:
    parts = {part.lower() for part in path.parts}
    for directory, doc_type in _DOC_TYPES.items():
        if directory in parts:
            return doc_type

    name = path.name.lower()
    for doc_type in _DOC_TYPES.values():
        if name.startswith(f"{doc_type}_"):
            return doc_type

    return "general"


def _extract_title(path: Path, content: str) -> str:
    lines = content.splitlines()
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# ") and stripped[2:].strip():
            return stripped[2:].strip()

    if path.suffix.lower() in {".md", ".html", ".htm", ".docx"}:
        return path.stem

    for line in lines:
        stripped = line.strip()
        if stripped:
            return stripped

    return path.stem


def _read_content(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix in {".html", ".htm"}:
        return _extract_html_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    raise ValueError(f"Unsupported document type: {suffix}")


def _split_front_matter(content: str) -> tuple[dict, str]:
    lines = content.splitlines()
    if not lines or lines[0].strip() != "---":
        return {}, content

    end_index = None
    for index, line in enumerate(lines[1:], start=1):
        if line.strip() == "---":
            end_index = index
            break
    if end_index is None:
        return {}, content

    metadata = {}
    for line in lines[1:end_index]:
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        parsed_value = value.strip()
        if parsed_value.lower() == "true":
            parsed_value = True
        elif parsed_value.lower() == "false":
            parsed_value = False
        metadata[key.strip()] = parsed_value
    return metadata, "\n".join(lines[end_index + 1 :]).strip()


def _raw_format(path: Path) -> str:
    suffix = path.suffix.lower().lstrip(".")
    return "html" if suffix == "htm" else suffix


def _normalize_lines(lines: Iterable[str]) -> str:
    cleaned = [line.strip() for line in lines if line and line.strip()]
    return "\n".join(cleaned)


def _extract_html_text(path: Path) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for node in soup(["script", "style", "noscript"]):
        node.extract()

    lines = []
    for element in soup.find_all(["h1", "h2", "h3", "p", "li", "td", "th"]):
        text = " ".join(element.get_text(" ", strip=True).split())
        if not text:
            continue
        if element.name == "h1":
            lines.append(f"# {text}")
        elif element.name in {"h2", "h3"}:
            lines.append(f"## {text}")
        else:
            lines.append(text)
    if not lines:
        lines = [" ".join(soup.get_text(" ", strip=True).split())]
    return _normalize_lines(lines)


def _extract_docx_text(path: Path) -> str:
    import docx

    document = docx.Document(path)
    lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif paragraph.style and paragraph.style.name.startswith("Heading"):
            lines.append(f"## {text}")
        else:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return _normalize_lines(lines)


def _extract_pdf_text(path: Path, max_pages: int = 8) -> str:
    try:
        return _extract_pdf_text_with_fitz(path, max_pages=max_pages)
    except Exception:
        return _extract_pdf_text_with_pypdf(path, max_pages=max_pages)


def _extract_pdf_text_with_fitz(path: Path, max_pages: int = 8) -> str:
    import fitz

    lines = []
    with fitz.open(path) as document:
        for page in document[:max_pages]:
            lines.extend((page.get_text("text") or "").splitlines())
    return _normalize_lines(lines)


def _extract_pdf_text_with_pypdf(path: Path, max_pages: int = 8) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    lines = []
    for page in reader.pages[:max_pages]:
        text = page.extract_text() or ""
        lines.extend(text.splitlines())
    return _normalize_lines(lines)
